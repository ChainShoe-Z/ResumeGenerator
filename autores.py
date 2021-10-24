
from docx2pdf import convert
from docx import Document
import docx
import tkinter as tk
import os
from PyPDF2 import PdfFileMerger

#set a TK GUI 
window = tk.Tk()
window.title('Auto Resume Generator')
window.geometry('500x300')


 #Entries
l = tk.Label(window, text='Company Name?', font=('Arial', 12), width=20, height=1)
e = tk.Entry(window, show = None, width=50)#display as plaintext
l.pack()
e.pack()

l2 = tk.Label(window, text='Company Address?', font=('Arial', 12), width=20, height=1)
e2 = tk.Entry(window, show = None, width=50)#display as plaintext
l2.pack()
e2.pack()

l3 = tk.Label(window, text='Position?', font=('Arial', 12), width=20, height=1)
e3 = tk.Entry(window, show = None, width=50)#display as plaintext
l3.pack()
e3.pack()

# 第5步，定义两个触发事件时的函数insert_point和insert_end（注意：因为Python的执行顺序是从上往下，所以函数一定要放在按钮的上面）
def insert_point(): # 在鼠标焦点处插入输入内容
    global var1 #set the input to global, to use them in later modification
    var1 = e.get()
    global var2 
    var2= e2.get()
    global var3
    var3 = e3.get()
    t.insert('insert', var1)
    t.insert('insert', var2)
    t.insert('insert', var3)

 
# 第6步，创建并放置两个按钮分别触发两种情况
b1 = tk.Button(window, text='MODIFY', width=10, height=2, command=insert_point)
b1.pack()

 
# 第7步，创建并放置一个多行文本框text用以显示，指定height=3为文本框是三个字符高度
t = tk.Text(window, height=3)
t.pack()
window.mainloop()


print(var1)
print(var2)
print(var3)
addr = var2.split(',',1)
print (addr)
street = addr[0]
provZip = addr[1]
print(street)

prov = provZip[:-7]
prov = prov[1:]
print(prov)

zipcode = provZip[-7:]
print(zipcode)



doc = docx.Document('cov.docx')

#modify here
#modify company name
doc.paragraphs[7].text = var1

#modify street
doc.paragraphs[8].text = street

#modify province
doc.paragraphs[9].text = prov

#modify zip code
doc.paragraphs[10].text = zipcode

#modify content in paragraph
doc.paragraphs[14].runs[6].text = var1
doc.paragraphs[14].runs[10].text = var3 + '.'

filename = var1 + 'Cov' +'.docx'
fPDF = var1 + 'Cov' + '.pdf'
doc.save(filename)

convert(filename, fPDF) #This will convert the excel file to pdf format, and save it as "this PC-->documents"


target_path = 'D:\programming\pyFiles'
pdf_lst = [f for f in os.listdir(target_path) if f.endswith('.pdf')]
pdf_lst = [os.path.join(target_path, filename) for filename in pdf_lst]

file_merger = PdfFileMerger()
for pdf in pdf_lst:
    file_merger.append(pdf)     # 合并pdf文件

file_merger.write("D:\programming\pyFiles/CV_ChengxuZhang.pdf")



